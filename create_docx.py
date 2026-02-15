
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_methodology_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('第三部分：核心方法 (Methodology) - PPT 内容策划', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('这部分是学术汇报的“硬核”环节，需要清晰展示理论依据、公式推导以及设计合理性。')

    # --- Page 5 ---
    doc.add_heading('Page 5: 理论基础：分位回归强化学习 (Theoretical Foundation)', level=1)
    
    doc.add_heading('标题: From Expected RL to Distributional RL', level=2)
    
    p = doc.add_paragraph()
    p.add_run('出发点:\n').bold = True
    p.add_run('在强随机性环境（Dynamic Currents & Obstacles）中，状态回报的期望值 (Expectation) 掩盖了潜在的风险。\n')
    p.add_run('例如：动作 A 有 50% 概率撞车 (-100)，50% 概率通过 (+10)，期望为 -45；动作 B 稳健绕行 (-5)。只看期望，Agent 可能会困惑。')

    p = doc.add_paragraph()
    p.add_run('关键公式与原理解析:\n').bold = True
    
    # Bellman Equation Explanation
    doc.add_heading('1. 传统 Bellman 方程 (Traditional Bellman Equation)', level=3)
    p = doc.add_paragraph()
    p.add_run('数学表达: ').bold = True
    p.add_run('Q(x, a) = E [ R(x, a) + γ Q(x\', a\') ]')
    
    p = doc.add_paragraph()
    p.add_run('公式解释: ').bold = True
    p.add_run('Bellman 方程是强化学习的核心，它描述了当前状态价值与未来状态价值之间的递归关系。\n')
    p.add_run('- Q(x, a): 在状态 x 执行动作 a 能获得的长期期望回报 (Scalar/数值)。\n')
    p.add_run('- E[...]: 期望算子，表示对所有可能的各种结果求平均。\n')
    p.add_run('- R(x, a): 即时奖励。\n')
    p.add_run('- γ (Gamma): 折扣因子，权衡如未来奖励的重要性。\n')
    p.add_run('- Q(x\', a\'): 下一个状态 x\' 采取最优动作 a\' 的期望价值。\n')
    p.add_run('核心含义: 传统的 Q-Learning 试图学习回报的“平均值”。但在高风险场景下，平均值会“拉平”极端的好坏情况，导致对危险不敏感。')

    # Distributional Bellman Equation
    doc.add_heading('2. 分布式 Bellman 方程 (Distributional Bellman Equation)', level=3)
    p = doc.add_paragraph()
    p.add_run('数学表达: ').bold = True
    p.add_run('Z(x, a) =D= R(x, a) + γ Z(x\', a\')')
    
    p = doc.add_paragraph()
    p.add_run('公式解释: ').bold = True
    p.add_run('这是 Distributional RL 对传统方程的推广。\n')
    p.add_run('- Z(x, a): 不再是一个数值，而是一个“随机变量” (Random Variable)，代表了回报的完整概率分布。\n')
    p.add_run('- =D=: 表示分布同分布 (Equal in Distribution)，即方程两边的概率分布是相同的。\n')
    p.add_run('核心含义: Agent 不仅预测平均能得多少分，还预测得分的波动范围（方差）、最坏可能（长尾风险）等。')

    p = doc.add_paragraph()
    p.add_run('可行性分析:\n').bold = True
    p.add_run('通过学习完整的分布 Z，Agent 能够知道某个动作“最坏情况”是什么，从而为避障提供信息基础。')

    # --- Page 6 ---
    doc.add_page_break()
    doc.add_heading('Page 6: 技术实现：隐式分位网络 (Implicit Quantile Networks, IQN)', level=1)
    
    doc.add_heading('标题: IQN Architecture for Navigation', level=2)
    
    p = doc.add_paragraph()
    p.add_run('核心思想:\n').bold = True
    p.add_run('将随机变量映射为分位函数 F^-1(τ)。不再输出固定的几个概率柱子，而是把概率 τ 作为输入，生成对应的回报值。')

    p = doc.add_paragraph()
    p.add_run('网络设计与公式:\n').bold = True
    
    doc.add_heading('1. 输入与嵌入', level=3)
    p = doc.add_paragraph('输入状态 x 和分位值 τ ∼ U([0, 1])。')
    p.add_run('\n分位嵌入 (Quantile Embedding):')
    p = doc.add_paragraph('ϕ(τ) := ReLU( Σ [cos(iπτ) * w_i + b_i] )')
    p.add_run('\n(利用余弦基函数将 0-1 之间的标量 τ 扩展为高维向量，捕捉不同概率水平的特征)')

    doc.add_heading('2. 输出分布近似', level=3)
    p = doc.add_paragraph('Z_τ(x, a) ≈ f( ψ(x) ⊙ ϕ(τ) )')
    p.add_run('\n(状态特征 ψ(x) 与 分位特征 ϕ(τ) 进行哈达玛积/逐元素相乘，融合环境信息与概率水平)')

    doc.add_heading('3. 训练目标: 分位回归损失 (Quantile Huber Loss)', level=3)
    p = doc.add_paragraph('L = (1/N) * Σ_i Σ_j [ ρ_τ_i ( Y_τ_j - Q(x, a, τ_i) ) ]')
    p.add_run('\n其中 ρ_τ 是分位回归的非对称损失函数，迫使网络预测准确的分位数。')

    p = doc.add_paragraph()
    p.add_run('可行性分析:\n').bold = True
    p.add_run('IQN 不需要预设分布的分段数，具有更高的样本利用率和拟合精度，适合连续变化的海洋环境风险建模。')

    # --- Page 7 ---
    doc.add_page_break()
    doc.add_heading('Page 7: 创新点：环境感知的自适应风险策略 (Adaptive Risk-Sensitive Policy)', level=1)
    
    doc.add_heading('标题: Adaptive Risk Sensitivity via CVaR', level=2)
    
    p = doc.add_paragraph()
    p.add_run('痛点:\n').bold = True
    p.add_run('固定风险偏好 (如一直保守) 导致导航效率低下；一直激进 (Risk-Neutral) 导致碰撞率高。')
    
    p = doc.add_paragraph()
    p.add_run('解决方案:\n').bold = True
    p.add_run('将条件风险价值 (CVaR) 作为决策准则，并基于环境动态调整置信水平 α。')

    p = doc.add_paragraph()
    p.add_run('关键公式:\n').bold = True
    
    doc.add_heading('1. 风险度量 (Risk Metric: CVaR)', level=3)
    p = doc.add_paragraph('Q_CVaR(x, a; α) = (1/α) * ∫[0 to α] F^-1_Z(τ) dτ')
    p.add_run('\n- α → 0: 极端风险厌恶 (只关注分布中最差的那一部分回报)。')
    p.add_run('\n- α → 1: 风险中性 (关注整体平均值，退化为普通 Q-Learning)。')

    doc.add_heading('2. 自适应机制 (Adaptive Mechanism)', level=3)
    p = doc.add_paragraph('α_adaptive = f(d_obs) = clip( d_min / D_safe, ε, 1.0 )')
    p.add_run('\n- d_min: 实时感知的最近障碍物距离。')
    p.add_run('\n- D_safe: 安全阈值 (e.g., 10m)。')
    p.add_run('\n- 含义: 距离越近，α 越小，策略越保守。')

    doc.add_heading('3. 最终决策逻辑', level=3)
    p = doc.add_paragraph('a* = argmax_a Q_CVaR(x, a; α_adaptive)')

    p = doc.add_paragraph()
    p.add_run('可行性分析:\n').bold = True
    p.add_run('- 在开阔水域 (α=1)，利用全分布信息做最高效规划。')
    p.add_run('\n- 在受限水域 (α << 1)，忽略高回报可能性，强制关注避险，从数学上保证了策略向“安全性”倾斜。')
    
    doc.save('methodology_presentation.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_methodology_doc()
