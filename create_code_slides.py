
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_structure_code_doc():
    doc = docx.Document()
    
    # --- Page 1: Algorithm Structure & Network Code ---
    doc.add_heading('Page 8: 算法结构详解：网络模型 (Algorithm Structure)', level=1)
    
    doc.add_heading('A. 神经网络架构 (ObsEncoder Implementation)', level=2)
    p = doc.add_paragraph('我们的网络设计采用了“分流感知-特征融合”的架构，代码位于 ')
    p.add_run('model.py class ObsEncoder').bold = True
    p.add_run('。')

    # Code Snippet 1: Forward Pass & Embedding
    doc.add_heading('1. 分位嵌入与前向传播 (Quantile Embedding & Forward)', level=3)
    p = doc.add_paragraph()
    p.add_run('核心代码解释:').bold = True
    p.add_run('\n这段代码展示了如何将标量风险因子 τ (taus) 通过余弦变换映射为高维特征，并注入到网络中。')
    
    code_snippet1 = """
# thirdparty/IQN/model.py

def calc_cos(self, batch_size, n_tau=8, cvar=1.0):
    # 分位采样与CVaR失真处理
    taus = torch.rand(batch_size, n_tau).to(self.device).unsqueeze(-1)
    taus = taus * cvar  # 关键：根据CVaR调整采样范围

    # 余弦嵌入 (Cosine Embedding)
    cos = torch.cos(taus * self.pis)
    return cos, taus

def forward(self, inputs, num_tau=8, cvar=1.0):
    # ... (前部分是对Sensors, Velocity, Goal的独立编码) ...
    
    # 1. 状态特征编码 (State Feature)
    features = torch.cat((v_features, g_features, s_featrues), 1)

    # 2. 分位特征编码 (Quantile Feature)
    cos, taus = self.calc_cos(batch_size, num_tau, cvar)
    cos_features = torch.relu(self.cos_embedding(cos)).view(batch_size, num_tau, 208)

    # 3. 特征融合 (Element-wise Fusion)
    # 状态特征与分位特征进行逐元素乘法
    features = (features.unsqueeze(1) * cos_features).view(batch_size * num_tau, 208)
    
    # ... (后续全连接层输出动作分布) ...
"""
    t = doc.add_paragraph(code_snippet1)
    t.style = 'No Spacing'
    for run in t.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_page_break()

    # --- Page 2: Adaptive Risk Mechanism ---
    doc.add_heading('Page 9: 核心创新：自适应风险调节 (Adaptive Risk Mechanism)', level=1)
    
    doc.add_heading('B. 动态 CVaR 调整逻辑 (Dynamic CVaR Logic)', level=2)
    p = doc.add_paragraph('代码位于 ')
    p.add_run('agent.py method adjust_cvar').bold = True
    p.add_run('。该函数实现了“距离越近，策略越保守”的核心思想。')

    # Code Snippet 2: Adjust CVaR
    code_snippet2 = """
# thirdparty/IQN/agent.py

def adjust_cvar(self, state):
    # 从状态向量提取声呐数据 (后22维)
    sonar_points = state[4:]
    
    # 计算最近障碍物距离 (Minimum Distance)
    closest_d = np.inf
    for i in range(0, len(sonar_points), 2):
        # ... (坐标处理与距离计算) ...
        closest_d = min(closest_d, np.linalg.norm(sonar_points[i:i+2]))
    
    # 核心自适应公式
    cvar = 1.0  # 默认：风险中性 (Risk Neutral)
    if closest_d < 10.0:
        # 当距离小于安全阈值(10m)时，降低 CVaR 值
        # CVaR 越小 = 越关注分布左尾部 = 越保守
        cvar = closest_d / 10.0

    return cvar
"""
    t = doc.add_paragraph(code_snippet2)
    t.style = 'No Spacing'
    for run in t.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    p = doc.add_paragraph('\n解释:')
    p.add_run('该机制实时计算最近障碍物。若距离充足 ($ \ge 10m $)，$CVaR=1.0$，模型使用全分布均值做决策（追求效率）。若距离迫近，CVaR 线性下降，迫使模型仅考虑“最坏情况”的回报，从而触发避障。')

    doc.add_page_break()

    # --- Page 3: Training Objective ---
    doc.add_heading('Page 10: 训练目标与损失 (Training Objective & Loss)', level=1)
    
    doc.add_heading('C. 分位回归损失 (Quantile Huber Loss)', level=2)
    p = doc.add_paragraph('为了训练网络准确预测回报分布的分位数，我们使用了 Quantile Huber Loss。代码位于 ')
    p.add_run('agent.py method train').bold = True
    p.add_run('。')

    # Code Snippet 3: Loss Calculation
    code_snippet3 = """
# thirdparty/IQN/agent.py (train method snippet)

# 1. 计算目标分布 (Target Distribution)
Q_targets_next, _ = self.qnetwork_target(next_states)
Q_targets_next = Q_targets_next.detach().max(2)[0].unsqueeze(1) 
Q_targets = rewards + (gamma * Q_targets_next * (1. - dones))

# 2. 计算预测分布 (Predicted Distribution)
Q_expected, taus = self.qnetwork_local(states)
# 提取对应动作的分布
Q_expected = Q_expected.gather(2, actions....) 

# 3. 计算 TD Error 矩阵
td_error = Q_targets - Q_expected

# 4. 分位回归损失 (关键部分)
# Huber Loss 部分：处理绝对误差
huber_l = calculate_huber_loss(td_error, 1.0)
# Quantile Weight 部分：非对称加权
# 核心思想：如果预测偏低(td_error<0)，用 tau 加权；反之用 (1-tau) 加权
quantil_l = abs(taus - (td_error.detach() < 0).float()) * huber_l

loss = quantil_l.mean()
"""
    t = doc.add_paragraph(code_snippet3)
    t.style = 'No Spacing'
    for run in t.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    p = doc.add_paragraph('\n总结:')
    p.add_run('这种非对称的损失函数迫使网络学习到真实的回报分布形状，而不仅仅是均值。这是 Distributional RL 能够在复杂海洋环境中捕捉潜在风险的数学根本。')

    doc.save('methodology_code_slides.docx')
    print("Code slides created successfully.")

if __name__ == "__main__":
    create_structure_code_doc()
